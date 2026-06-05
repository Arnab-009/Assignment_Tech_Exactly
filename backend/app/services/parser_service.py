"""Document parsing service.

Extracts plain text **and lightweight structure** (page count, word count and
tables) from PDF, DOCX and TXT byte streams. Tables are a first-class citizen:
they are flattened into a compact textual form so the LLM can reason about
tabular data instead of ignoring it.

The service never raises on bad input — a corrupt or encrypted file yields a
``ParseResult`` with ``status="error"`` and an explanatory note, so one bad
file can never abort a whole batch.
"""
from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field
from typing import Optional

import fitz  # PyMuPDF — fast text extraction + page count
import pdfplumber  # robust table extraction
from docx import Document

logger = logging.getLogger(__name__)

# Sentinel values returned in place of extracted text.
EMPTY_DOCUMENT = "[EMPTY_DOCUMENT]"
PARSE_ERROR = "[PARSE_ERROR]"

PDF_MIME = "application/pdf"
DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
TXT_MIME = "text/plain"

# Keep flattened tables bounded so they never dominate the LLM context window.
_MAX_TABLES_CHARS = 8_000


@dataclass
class ParseResult:
    """Outcome of parsing a single document."""

    text: str
    status: str  # "ok" | "empty" | "error"
    page_count: Optional[int] = None
    word_count: int = 0
    table_count: int = 0
    tables_text: str = ""
    note: Optional[str] = None


class ParserService:
    """Extract clean text + structure from supported document byte streams."""

    def __init__(self, max_chars: int = 50_000) -> None:
        self._max_chars = max_chars

    # -- Public API ---------------------------------------------------------
    def parse(self, file_bytes: bytes, mime_type: str) -> ParseResult:
        """Parse a document into text + structural metadata.

        Never raises: returns a ``ParseResult`` whose ``status`` describes the
        outcome (``ok`` / ``empty`` / ``error``).
        """
        if not file_bytes:
            return ParseResult(text=EMPTY_DOCUMENT, status="empty")

        normalized = (mime_type or "").split(";")[0].strip().lower()
        try:
            if normalized == PDF_MIME:
                result = self._parse_pdf(file_bytes)
            elif normalized == DOCX_MIME:
                result = self._parse_docx(file_bytes)
            elif normalized == TXT_MIME or normalized.startswith("text/"):
                result = self._parse_txt(file_bytes)
            else:
                logger.warning("Unsupported mime type for parsing: %s", mime_type)
                return ParseResult(
                    text=PARSE_ERROR,
                    status="error",
                    note=f"Unsupported document type: {mime_type}",
                )
        except Exception as exc:  # noqa: BLE001 - parser must never crash a batch
            logger.exception("Failed to parse document (mime=%s)", mime_type)
            return ParseResult(
                text=PARSE_ERROR, status="error", note=f"Parse error: {exc}"
            )

        result.text = (result.text or "").strip()
        if not result.text:
            result.status = "empty"
            result.text = EMPTY_DOCUMENT
            result.word_count = 0
            return result

        result.word_count = len(result.text.split())
        result.text = result.text[: self._max_chars]
        result.tables_text = result.tables_text[:_MAX_TABLES_CHARS]
        return result

    def extract_text(self, file_bytes: bytes, mime_type: str) -> str:
        """Backwards-compatible helper returning only the extracted text."""
        return self.parse(file_bytes, mime_type).text

    # -- Per-format parsers -------------------------------------------------
    def _parse_pdf(self, file_bytes: bytes) -> ParseResult:
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            if doc.needs_pass:
                # Best-effort: try an empty password before giving up.
                if not doc.authenticate(""):
                    return ParseResult(
                        text=PARSE_ERROR,
                        status="error",
                        note="PDF is password-protected / encrypted.",
                    )
            page_count = doc.page_count
            text = "\n".join(page.get_text("text") for page in doc)

        tables_text, table_count = self._extract_pdf_tables(file_bytes)
        return ParseResult(
            text=text,
            status="ok",
            page_count=page_count,
            table_count=table_count,
            tables_text=tables_text,
        )

    def _parse_docx(self, file_bytes: bytes) -> ParseResult:
        document = Document(io.BytesIO(file_bytes))
        parts = [para.text for para in document.paragraphs if para.text.strip()]

        table_blocks: list[str] = []
        for index, table in enumerate(document.tables, start=1):
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            block = self._format_table(rows, index)
            if block:
                table_blocks.append(block)

        # Append a readable version of the tables to the body text too, so even
        # if structured table parsing is imperfect the content is not lost.
        text = "\n".join(parts)
        if table_blocks:
            text = f"{text}\n\n{chr(10).join(table_blocks)}"

        return ParseResult(
            text=text,
            status="ok",
            page_count=None,
            table_count=len(table_blocks),
            tables_text="\n\n".join(table_blocks),
        )

    def _parse_txt(self, file_bytes: bytes) -> ParseResult:
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text = file_bytes.decode("latin-1", errors="replace")
        return ParseResult(text=text, status="ok")

    # -- Table helpers ------------------------------------------------------
    def _extract_pdf_tables(self, file_bytes: bytes) -> tuple[str, int]:
        """Extract tables from a PDF with pdfplumber. Best-effort, never raises."""
        blocks: list[str] = []
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                counter = 0
                for page in pdf.pages:
                    for raw in page.extract_tables() or []:
                        counter += 1
                        block = self._format_table(raw, counter)
                        if block:
                            blocks.append(block)
                    if sum(len(b) for b in blocks) > _MAX_TABLES_CHARS:
                        break
        except Exception as exc:  # noqa: BLE001 - tables are a bonus, not critical
            logger.info("PDF table extraction skipped: %s", exc)
            return "", 0
        return "\n\n".join(blocks), len(blocks)

    @staticmethod
    def _format_table(rows: list[list], index: int) -> str:
        """Flatten a table (list of rows of cells) into compact pipe-text."""
        cleaned: list[str] = []
        for row in rows:
            cells = [
                " ".join(str(cell).split()) if cell is not None else ""
                for cell in row
            ]
            if any(cells):
                cleaned.append(" | ".join(cells))
        if not cleaned:
            return ""
        return f"[Table {index}]\n" + "\n".join(cleaned)
