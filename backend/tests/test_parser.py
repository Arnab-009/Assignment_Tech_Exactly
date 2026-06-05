"""Tests for app.services.parser_service."""
from __future__ import annotations

import io
from pathlib import Path

import pytest

from app.services.parser_service import (
    DOCX_MIME,
    EMPTY_DOCUMENT,
    PARSE_ERROR,
    PDF_MIME,
    TXT_MIME,
    ParserService,
)


@pytest.fixture()
def parser() -> ParserService:
    return ParserService(max_chars=50_000)


# ── TXT parsing ──────────────────────────────────────────────────────────────
class TestParseTxt:
    def test_extracts_utf8_text(self, parser: ParserService, sample_txt_bytes: bytes):
        result = parser.extract_text(sample_txt_bytes, TXT_MIME)
        assert "plain text test fixture" in result
        assert result != EMPTY_DOCUMENT
        assert result != PARSE_ERROR

    def test_handles_latin1_fallback(self, parser: ParserService):
        # Latin-1 encoded string that's invalid UTF-8
        latin1_bytes = "café résumé".encode("latin-1")
        result = parser.extract_text(latin1_bytes, TXT_MIME)
        assert "caf" in result
        assert result != PARSE_ERROR

    def test_generic_text_mime(self, parser: ParserService):
        """text/csv and similar text/* types should be handled as plain text."""
        data = b"col1,col2\nval1,val2"
        result = parser.extract_text(data, "text/csv")
        assert "col1" in result

    def test_empty_bytes_returns_sentinel(self, parser: ParserService):
        assert parser.extract_text(b"", TXT_MIME) == EMPTY_DOCUMENT

    def test_whitespace_only_returns_empty(self, parser: ParserService):
        assert parser.extract_text(b"   \n\t  ", TXT_MIME) == EMPTY_DOCUMENT


# ── Truncation ───────────────────────────────────────────────────────────────
class TestTruncation:
    def test_text_is_truncated_to_max_chars(self):
        parser = ParserService(max_chars=20)
        data = ("a" * 100).encode()
        result = parser.extract_text(data, TXT_MIME)
        assert len(result) == 20


# ── Error handling ───────────────────────────────────────────────────────────
class TestErrorHandling:
    def test_unsupported_mime_returns_parse_error(self, parser: ParserService):
        result = parser.extract_text(b"data", "application/zip")
        assert result == PARSE_ERROR

    def test_corrupt_pdf_returns_parse_error(self, parser: ParserService):
        result = parser.extract_text(b"not a pdf", PDF_MIME)
        assert result == PARSE_ERROR

    def test_corrupt_docx_returns_parse_error(self, parser: ParserService):
        result = parser.extract_text(b"not a docx", DOCX_MIME)
        assert result == PARSE_ERROR


# ── PDF parsing (requires a real tiny PDF) ───────────────────────────────────
class TestParsePdf:
    @pytest.fixture()
    def tiny_pdf_bytes(self) -> bytes:
        """Create a minimal valid PDF with PyMuPDF for testing."""
        import fitz

        doc = fitz.open()
        page = doc.new_page(width=200, height=200)
        page.insert_text((50, 100), "Hello PDF World", fontsize=12)
        data = doc.tobytes()
        doc.close()
        return data

    def test_extracts_text_from_pdf(self, parser: ParserService, tiny_pdf_bytes: bytes):
        result = parser.extract_text(tiny_pdf_bytes, PDF_MIME)
        assert "Hello PDF World" in result
        assert result != EMPTY_DOCUMENT


# ── DOCX parsing (requires a real tiny DOCX) ────────────────────────────────
class TestParseDocx:
    @pytest.fixture()
    def tiny_docx_bytes(self) -> bytes:
        """Create a minimal valid DOCX with python-docx for testing."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Hello DOCX World")
        doc.add_paragraph("Second paragraph with useful info.")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_extracts_text_from_docx(self, parser: ParserService, tiny_docx_bytes: bytes):
        result = parser.extract_text(tiny_docx_bytes, DOCX_MIME)
        assert "Hello DOCX World" in result
        assert "Second paragraph" in result

    @pytest.fixture()
    def docx_with_table_bytes(self) -> bytes:
        """DOCX containing a table to verify table extraction."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Document with table")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "Alpha"
        table.cell(1, 1).text = "100"
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_extracts_table_content(self, parser: ParserService, docx_with_table_bytes: bytes):
        result = parser.extract_text(docx_with_table_bytes, DOCX_MIME)
        assert "Name" in result
        assert "Alpha" in result
